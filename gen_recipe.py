import json
from docx import Document
import subprocess
import os

def generate_recipe_docx(json_file, file_name):
    # Load the recipe from the JSON file
    with open(json_file, 'r') as file:
        recipe = json.load(file)
    
    # File names with according file endings
    docx_file = file_name + ".docx"
    pages_file = file_name + ".pages"


    print(docx_file)
    # Create a new Word document
    doc = Document()
    
    # Add the title
    doc.add_heading(recipe['title'], level=1)
    
    # Add the ingredients
    doc.add_heading('Ingredients', level=2)
    for ingredient in recipe['ingredients']:
        doc.add_paragraph(f"{ingredient['amount']} - {ingredient['name']}")
    
    # Add the description (cooking steps)
    doc.add_heading('Instructions', level=2)
    for step in recipe['description']:
        doc.add_paragraph(step)
    
    # Save the document as a .docx file
    doc.save(docx_file)
    print(f"Recipe saved to {docx_file}")
    
    # Convert the .docx file to .pages using AppleScript
    convert_to_pages(docx_file, pages_file)

    os.remove(docx_file)

def convert_to_pages(docx_file, pages_file):
    applescript = f'''
    tell application "Pages"
        open POSIX file "/Users/laurent/Documents/Projects/Python/RecipeGenerator/{docx_file}"
        set theDoc to document 1
        save theDoc in POSIX file "/Users/laurent/Documents/Projects/Python/RecipeGenerator/{pages_file}"
        close theDoc
    end tell
    '''
    subprocess.run(["osascript", "-e", applescript])
    print(f"Converted {docx_file} to {pages_file}")

# Example usage
generate_recipe_docx("carbonara.json", "Carbonara")
